from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = r"C:\GestaoVersus\app32\output\investidores\Resumo_Executivo_Investidores_Recarga_Eletrica_v9_final.docx"


def set_base_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    style.font.size = Pt(10.5)
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


doc = Document()
set_base_style(doc)

add_title(doc, "Resumo Executivo para Investidores")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Infraestrutura de recarga elétrica | Versão final revisada com base no PDF 03").italic = True

doc.add_paragraph(
    "A operação foi desenhada para capturar a expansão estrutural da mobilidade elétrica no Brasil, combinando venda de kits, software de gestão e frentes complementares de implantação, assessoria, peças e manutenção. "
    "Na modelagem atual, apenas 2 produtos sustentam efetivamente a rampa financeira — kit/equipamento principal e software de gestão — o que torna a tese conservadora e preserva upside concreto pela ativação futura de mais 4 frentes do portfólio."
)

add_h1(doc, "1. A empresa")
doc.add_paragraph(
    "A estrutura de gestão prevê conselho de administração, diretoria e gerências comercial, operacional e adm/fin, com César Lisboa na condução estratégica. "
    "O capital é 100% próprio na largada, com execução faseada e foco em segurança do investimento, disciplina operacional e valorização da cadeia produtiva — colaboradores, fornecedores e clientes."
)

add_h1(doc, "2. Produtos e base econômica")
add_h2(doc, "Produtos que sustentam o resultado atual")
for text in [
    "Kit/equipamento principal: 300 unidades pós-ramp-up, preço unitário de R$ 80.000, MC% de 33,4%, MC unitária de R$ 26.736 e MC total de R$ 8.020.800.",
    "Software de gestão: base ativa de 4.000 contratos, preço unitário de R$ 1.600, MC% de 75,8%, MC unitária de R$ 1.212 e MC total de R$ 4.848.000.",
]:
    add_bullet(doc, text)

add_h2(doc, "Produtos adicionais passíveis de ativação")
for text in [
    "Implantação/instalação: economicamente modelado, porém sem volume e sem receita na rampa atual.",
    "Estudo de viabilidade/assessoria: ainda não ativado na modelagem.",
    "Peças e upgrades: ainda não ativado na modelagem.",
    "Manutenção especializada: ainda não ativado na modelagem.",
]:
    add_bullet(doc, text)

doc.add_paragraph(
    "Com isso, o faturamento mensal pós-ramp-up da versão atual soma R$ 30,4 milhões, com margem de contribuição mensal de R$ 12,8688 milhões. "
    "A leitura correta é de 2 produtos contribuindo para o resultado e 4 frentes adicionais passíveis de ativação ao longo da execução."
)

add_h1(doc, "3. Demanda de mercado e déficit de infraestrutura")
doc.add_paragraph(
    "Com 411.869 veículos plug-in em circulação e 21.061 pontos públicos e semipúblicos de recarga, o Brasil opera hoje com relação de 19,6 veículos por ponto. "
    "Adotando referência de 10 veículos por ponto, o déficit atual é de 20.126 pontos de recarga."
)

tbl = doc.add_table(rows=1, cols=5)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Ano"
hdr[1].text = "Frota projetada"
hdr[2].text = "Pontos necessários (10:1)"
hdr[3].text = "Déficit acumulado"
hdr[4].text = "Demanda incremental"
for row in [
    ("2026", "639.869", "63.987", "42.926", "22.800"),
    ("2027", "1.004.869", "100.487", "79.426", "36.500"),
    ("2028", "1.588.869", "158.887", "137.826", "58.400"),
    ("2029", "2.523.869", "252.387", "231.326", "93.500"),
    ("2030", "4.023.869", "402.387", "381.326", "150.000"),
]:
    cells = tbl.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

doc.add_paragraph(
    "As projeções de 2027 a 2030 foram tratadas como cenário-base de planejamento, alinhado à tendência setorial de expansão da frota. "
    "A mensagem para investidores é clara: o mercado já nasce com déficit relevante e tende a ampliar a pressão por infraestrutura ao longo do horizonte."
)

add_h1(doc, "4. Estrutura e custo fixo")
for text in [
    "Comercial: setup de R$ 340 mil e custo fixo mensal de R$ 110 mil.",
    "Operacional: setup de R$ 331 mil e custo fixo mensal de R$ 109 mil.",
    "Adm/Fin: setup de R$ 102,5 mil e custo fixo mensal de R$ 92,5 mil.",
]:
    add_bullet(doc, text)
doc.add_paragraph(
    "O setup total das frentes soma R$ 773,5 mil, com custo fixo mensal consolidado de R$ 311,5 mil."
)

add_h1(doc, "5. Resultado líquido em cenários de venda")
tbl2 = doc.add_table(rows=1, cols=6)
tbl2.style = "Table Grid"
hdr = tbl2.rows[0].cells
hdr[0].text = "Cenário"
hdr[1].text = "Receita mensal"
hdr[2].text = "MC mensal"
hdr[3].text = "Custo fixo mensal"
hdr[4].text = "Impostos"
hdr[5].text = "Resultado líquido mensal"
for row in [
    ("20%", "R$ 6.080.000", "R$ 2.573.760", "R$ 311.500", "R$ 837.036", "R$ 1.425.224"),
    ("50%", "R$ 15.200.000", "R$ 6.434.400", "R$ 311.500", "R$ 2.265.473", "R$ 3.857.427"),
    ("100%", "R$ 30.400.000", "R$ 12.868.800", "R$ 311.500", "R$ 4.646.201", "R$ 7.911.099"),
]:
    cells = tbl2.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

add_h1(doc, "6. Investimento necessário")
for text in [
    "Caixa / capital de giro: R$ 1,9 milhão.",
    "Estoques / containers: R$ 8,82 milhões.",
    "Investimentos fixos mapeados: R$ 462 mil.",
    "Total de investimento: R$ 11,5 milhões.",
]:
    add_bullet(doc, text)
doc.add_paragraph(
    "O modelo considerou taxa de dólar de R$ 5,00. O investimento segue concentrado em estoque e capacidade de entrega, com capital de giro suficiente para reduzir risco de execução no ramp-up."
)

add_h1(doc, "7. Premissas do investidor")
for text in [
    "Distribuição de 25% a partir de nov/2026; nesta versão, a métrica foi recalculada com o fluxo mensal do próprio plano.",
    "Custo de oportunidade: 12% a.a.",
    "TIR estimada: 142,2% a.a.",
    "VPL estimado: R$ 10.797.483.",
    "Payback estimado: ~10,2 meses.",
]:
    add_bullet(doc, text)

add_h1(doc, "8. Fluxo de caixa resumido")
doc.add_paragraph(
    "Negócio — 1 mês após o ramp-up (set/2027): receita de R$ 30,4 milhões, margem de contribuição de R$ 12,8688 milhões, custo fixo mensal de R$ 311,5 mil e resultado líquido de R$ 7,9111 milhões."
)
doc.add_paragraph(
    "Negócio — ano cheio de 2028: receita anual de R$ 364,8 milhões, margem de contribuição anualizada de R$ 154,4256 milhões, custo fixo anual de R$ 3,738 milhões e resultado líquido anualizado de R$ 94,9332 milhões."
)
doc.add_paragraph(
    "Investidor — aporte inicial de R$ 11,5 milhões em jul/2026, distribuição mensal estimada de R$ 1.977.774,75 em set/2027 e distribuição anualizada de R$ 23.733.297 no ano cheio de 2028."
)

add_h1(doc, "9. Fechamento")
doc.add_paragraph(
    "A tese fecha com 2 produtos contribuindo para o resultado atual e pode ganhar potência adicional com a ativação de mais 4 frentes do portfólio. "
    "Isso preserva uma leitura conservadora do caso base e mantém upside concreto de receita, margem, recorrência e captura de valor, sem depender dessas frentes adicionais para viabilizar o investimento."
)
doc.add_paragraph(
    "Em síntese, trata-se de uma operação posicionada em mercado com déficit estrutural de infraestrutura, base recorrente crescente, investimento integralmente mapeado e retorno potencial relevante para o investidor."
)

doc.save(OUT)
print(OUT)
